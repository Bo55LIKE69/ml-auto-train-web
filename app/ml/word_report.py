# -*- coding: utf-8 -*-
"""
Word 实验报告生成器（规格书 §5.5 价值锚点）。

基于 metrics.json + profile 信息 + charts/*.png 生成中文实验报告，
七章固定骨架：
    第一章 实验概述
    第二章 数据说明
    第三章 实验设置
    第四章 模型对比
    第五章 最优模型分析
    第六章 局限与改进建议
    第七章 附录

样式规范：
- 正文：宋体/Times New Roman，小四号（12pt），1.5 倍行距
- 标题：黑体，三号（16pt）一级标题，四号（14pt）二级标题
- 表格：三线表（顶线、栏目线、底线），表头加粗
- 图片：居中，下方标注"图 X-X XXX"，150dpi
- 每章底部小字免责声明
"""
import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.config import REPORT_CHART_DPI, REPORT_DISCLAIMER

# 中文字体映射（docx 需同时设置 ascii 与 eastAsia）
BODY_FONT = "宋体"
HEADING_FONT = "黑体"
ASCII_FONT = "Times New Roman"


def _set_run_font(run, font_name=BODY_FONT, size=12, bold=False, color=None):
    """统一设置 run 的字体：中文字体 + ASCII 字体 + 字号 + 粗细。"""
    run.font.name = ASCII_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_heading_styled(doc, text, level=1):
    """统一标题样式：黑体，16pt 一级 / 14pt 二级。"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        _set_run_font(run, HEADING_FONT, 16 if level == 1 else 14, bold=True)
    return heading


def _add_body(doc, text, size=12, bold=False, align=None):
    """正文段落：宋体小四（12pt），1.5 倍行距。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, BODY_FONT, size, bold)
    p.paragraph_format.line_spacing = 1.5
    if align:
        p.alignment = align
    return p


def _set_cell_borders(cell, top=None, bottom=None):
    """设置单元格边框（三线表：顶线/底线粗）。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge, val in (("top", top), ("bottom", bottom)):
        if val is None:
            continue
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(val))
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tc_pr.append(borders)


def _add_three_line_table(doc, headers, data):
    """
    添加三线表（学术论文标准格式）：
    顶线粗、栏目线细、底线粗，表头加粗。
    """
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        _set_run_font(run, BODY_FONT, 10.5, bold=True)
        _set_cell_borders(cell, top=12, bottom=6)

    # 数据行
    for r_idx, row_data in enumerate(data):
        is_last = r_idx == len(data) - 1
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            _set_run_font(run, BODY_FONT, 10.5)
            if is_last:
                _set_cell_borders(cell, bottom=12)
    return table


def _add_image_caption(doc, caption):
    """图片下方标注：图 X-X XXX（居中，小五号）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    _set_run_font(run, BODY_FONT, 9)
    return p


def _add_disclaimer(doc):
    """章节底部小字免责声明。"""
    p = doc.add_paragraph()
    run = p.add_run(REPORT_DISCLAIMER)
    _set_run_font(run, BODY_FONT, 8)
    p.paragraph_format.line_spacing = 1.0
    return p


def _metric_display_name(key: str) -> str:
    """指标 key → 中文名（规格书附录 B）。"""
    names = {
        "accuracy": "准确率", "precision": "精确率", "recall": "召回率",
        "f1": "F1", "auc": "AUC", "kappa": "Kappa", "mcc": "MCC",
        "r2": "R²", "mae": "MAE", "mse": "MSE", "rmse": "RMSE",
        "train_time_s": "训练耗时(s)",
    }
    return names.get(key, key)


def generate_word_report(result, df, y, meta, save_path):
    """
    生成七章 Word 实验报告。
    result: run_pipeline 返回值；df/y/meta: 预处理阶段元信息。
    """
    doc = Document()
    task = "分类" if result["task_type"] == "classification" else "回归"
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    target = result["target_col"]
    best = result["best_model"]
    out_dir = save_path.parent

    # ============ 第一章 实验概述 ============
    _add_heading_styled(doc, "第一章 实验概述", 1)
    _add_body(doc, f"本实验使用「AutoML 毕设实验工作台」对数据集进行机器学习建模分析，"
                   f"目标标签为「{target}」，任务类型为「{task}」。")
    _add_body(doc, f"实验名称：{target} 预测实验")
    _add_body(doc, f"实验日期：{now}")
    _add_body(doc, f"数据集来源：{result.get('source_file', meta.get('source_file', '用户上传'))}")
    _add_body(doc, f"作者：（请填写）")
    _add_body(doc, f"工具版本：AutoML 毕设实验工作台 v1.0")
    _add_body(doc, f"样本数量：{result['n_samples']} 条，特征数量：{result['n_features_raw']} 个"
                   f"（预处理后 {result['n_features_after_prep']} 维）。")
    _add_disclaimer(doc)

    # ============ 第二章 数据说明 ============
    _add_heading_styled(doc, "第二章 数据说明", 1)
    _add_heading_styled(doc, "2.1 数据规模", 2)
    _add_three_line_table(doc,
        ["指标", "数值"],
        [["样本数量", str(result["n_samples"])],
         ["原始特征数", str(result["n_features_raw"])],
         ["预处理后特征数", str(result["n_features_after_prep"])],
         ["目标列", target],
         ["任务类型", task],
         ["类别数", str(len(result["class_names"])) if result["class_names"] else "连续数值"],
         ["类别分布", " / ".join(result["class_names"]) if result["class_names"] else "—"]])

    _add_heading_styled(doc, "2.2 数据质量问题", 2)
    if result["warnings"]:
        for w in result["warnings"]:
            _add_body(doc, f"- {w}")
    else:
        _add_body(doc, "- 无显著数据质量问题")
    _add_heading_styled(doc, "2.3 剔除列说明", 2)
    if result["drop_cols"]:
        _add_body(doc, "已剔除以下无意义列：" + "、".join(result["drop_cols"]))
    else:
        _add_body(doc, "无剔除列。")
    _add_disclaimer(doc)

    # ============ 第三章 实验设置 ============
    _add_heading_styled(doc, "第三章 实验设置", 1)
    _add_heading_styled(doc, "3.1 任务定义", 2)
    _add_body(doc, f"本实验为{task}任务，预测目标为「{target}」。")
    _add_heading_styled(doc, "3.2 数据划分策略", 2)
    _add_body(doc, "使用 7:3 比例划分训练集与测试集（随机种子 42，保证可复现）。"
                   "分类任务采用分层抽样（Stratified Split），保证训练/测试集类别比例一致。")
    _add_heading_styled(doc, "3.3 预处理流水线", 2)
    _add_body(doc, "数值列：缺失值以中位数填充，随后标准化（StandardScaler）。")
    _add_body(doc, "类别列：缺失值以众数填充，随后独热编码（OneHotEncoder）。")
    _add_body(doc, "预处理管道仅在训练集上拟合，测试集只做变换，避免数据泄漏。")
    _add_heading_styled(doc, "3.4 评价指标", 2)
    if task == "分类":
        _add_three_line_table(doc,
            ["指标", "含义", "越大越好？"],
            [["Accuracy", "准确率", "是"],
             ["Precision", "精确率（macro）", "是"],
             ["Recall", "召回率（macro）", "是"],
             ["F1", "精确率与召回率的调和平均（macro）", "是"],
             ["AUC", "ROC 曲线下面积（二分类）", "是"],
             ["Kappa", "Cohen's Kappa 一致性系数", "是"],
             ["MCC", "Matthews 相关系数", "是"]])
    else:
        _add_three_line_table(doc,
            ["指标", "含义", "越小越好？"],
            [["MAE", "平均绝对误差", "是"],
             ["MSE", "均方误差", "是"],
             ["RMSE", "均方根误差", "是"],
             ["R²", "决定系数", "否（越大越好）"]])
    _add_disclaimer(doc)

    # ============ 第四章 模型对比 ============
    _add_heading_styled(doc, "第四章 模型对比", 1)
    _add_heading_styled(doc, "4.1 模型排名表", 2)
    _add_body(doc, f"共训练 {len(result['models'])} 个模型，按 "
                   f"{'F1(macro)' if task == '分类' else 'R²'} 排序如下（三线表格式，可直接用于论文）：")
    metric_keys = list(result["models"][0]["metrics"].keys())
    headers = ["排名", "模型"] + [_metric_display_name(k) for k in metric_keys]
    rows = []
    for rank, m in enumerate(sorted(
            result["models"], key=lambda x: x["metrics"]["f1"] if task == "分类" else x["metrics"]["r2"], reverse=True), 1):
        rows.append([str(rank), m["name"]] + [str(m["metrics"].get(k, "")) for k in metric_keys])
    _add_three_line_table(doc, headers, rows)

    _add_heading_styled(doc, "4.2 指标对比图", 2)
    p = out_dir / "metrics_comparison.png"
    if p.exists():
        doc.add_picture(str(p), width=Cm(14))
        _add_image_caption(doc, "图 4-1 模型指标对比（测试集）")
    _add_disclaimer(doc)

    # ============ 第五章 最优模型分析 ============
    _add_heading_styled(doc, "第五章 最优模型分析", 1)
    _add_heading_styled(doc, "5.1 最优模型", 2)
    _add_body(doc, f"最优模型为「{best['name']}」（依据：{best['reason']}）。")
    metric_str = "，".join(f"{_metric_display_name(k)}={v}" for k, v in best["metrics"].items() if v is not None)
    _add_body(doc, f"测试集指标：{metric_str}")

    _add_heading_styled(doc, "5.2 混淆矩阵", 2)
    p = out_dir / "confusion_matrix.png"
    if p.exists():
        doc.add_picture(str(p), width=Cm(12))
        _add_image_caption(doc, "图 5-1 混淆矩阵（最优模型 · 测试集）")

    _add_heading_styled(doc, "5.3 特征重要性", 2)
    p = out_dir / "feature_importance.png"
    if p.exists():
        doc.add_picture(str(p), width=Cm(12))
        _add_image_caption(doc, "图 5-2 特征重要性 Top-15")
    if result["feature_importance"]:
        imp_rows = [[str(i), item["feature"], str(item["importance"])]
                    for i, item in enumerate(result["feature_importance"][:10], 1)]
        _add_three_line_table(doc, ["排名", "特征", "重要性"], imp_rows)
    _add_disclaimer(doc)

    # ============ 第六章 局限与改进建议 ============
    _add_heading_styled(doc, "第六章 局限与改进建议", 1)
    _add_body(doc, f"- 数据量有限（{result['n_samples']} 条），模型性能存在偶然性，建议扩充样本。")
    _add_body(doc, "- 仅使用模型默认参数，未进行网格搜索调优，后续可对最优模型做 GridSearchCV 超参数优化。")
    _add_body(doc, "- 类别列采用独热编码，类别数较多时特征维度膨胀，可尝试目标编码（Target Encoding）。")
    _add_body(doc, "- 未处理类别不平衡问题，若样本分布不均，可尝试 SMOTE 过采样或调整类别权重。")
    if result["class_names"] and len(result["class_names"]) > 2:
        _add_body(doc, f"- 本任务为多分类（{len(result['class_names'])} 类），AUC 指标仅在二分类时有效，"
                       f"多分类场景建议以 F1(macro) 为主要参考。")
    _add_body(doc, "- 后续可引入交叉验证（K-Fold）进一步验证模型稳定性，当前为单次 7:3 划分。")
    _add_disclaimer(doc)

    # ============ 第七章 附录 ============
    _add_heading_styled(doc, "第七章 附录", 1)
    _add_heading_styled(doc, "7.1 环境清单", 2)
    try:
        import platform
        import sklearn
        import pandas as pd
        import numpy as np
        env_rows = [
            ["Python", platform.python_version()],
            ["scikit-learn", sklearn.__version__],
            ["pandas", pd.__version__],
            ["numpy", np.__version__],
        ]
        _add_three_line_table(doc, ["组件", "版本"], env_rows)
    except Exception:
        pass
    _add_heading_styled(doc, "7.2 完整指标数据", 2)
    full_rows = []
    for m in result["models"]:
        full_rows.append([m["name"]] + [str(m["metrics"].get(k, "")) for k in metric_keys])
    _add_three_line_table(doc, ["模型"] + [_metric_display_name(k) for k in metric_keys], full_rows)
    _add_disclaimer(doc)

    doc.save(str(save_path))
    return save_path
