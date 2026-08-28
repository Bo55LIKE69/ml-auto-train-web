# -*- coding: utf-8 -*-
"""验证 Word 报告与 pipeline_ir 完整性。"""
import io
import sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
sys.path.insert(0, ".")

import json
from docx import Document
from pathlib import Path

out = Path(r"D:\ML_help\outputs\verify_v2")

# 1. Word 报告
doc = Document(str(out / "report.docx"))
paras = [p.text for p in doc.paragraphs if p.text.strip()]
print("Word 段落数:", len(paras))
print("前 12 段:")
for p in paras[:12]:
    print("  ", p[:60])
print("表格数:", len(doc.tables))
for i, t in enumerate(doc.tables[:3]):
    print(f"  表{i+1}: {len(t.rows)}行 x {len(t.columns)}列, 表头={[c.text for c in t.rows[0].cells]}")

# 2. pipeline_ir
ir = json.loads((out / "pipeline_ir.json").read_text(encoding="utf-8"))
print("\nIR data_info:", json.dumps(ir["data_info"], ensure_ascii=False)[:200])
print("IR setup_config.fold:", ir["setup_config"]["fold"], "sort:", ir["setup_config"]["sort_metric"])
print("IR model_results.best:", ir["model_results"]["best_model_name"])
print("IR env:", ir["environment"]["sklearn_version"])
